import streamlit as st
import pandas as pd
import scipy.stats as stats
import seaborn as sns
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D
from scikit_posthocs import posthoc_dunn
from statannotations.Annotator import Annotator
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime
from lifelines import KaplanMeierFitter
from lifelines.statistics import logrank_test

# --- Color Palettes ---
PALETTES = {
    "🔴 红蓝经典": ['#E64B35', '#4DBBD5', '#00A087', '#3C5488', '#F39B7F', '#8491B4', '#91D1C2', '#DC0000', '#7E6148', '#B09C85'],
    "🔵 冷色专业": ['#3B4992', '#EE0000', '#008B45', '#631879', '#008280', '#BB0021', '#5F559B', '#A20056', '#808180', '#1B1919'],
    "🟠 暖色医学": ['#BC3C29', '#0072B5', '#E18727', '#20854E', '#7876B1', '#6F99AD', '#FFDC91', '#EE4C97'],
    "🟢 柔和自然": "Set2",
    "🟣 高对比": "Set1",
    "🔵 蓝色渐变": "Blues",
    "🤎 色盲友好": ['#0077BB', '#33BBEE', '#009988', '#EE7733', '#CC3311', '#EE3377', '#BBBBBB'],
    "⚫ 灰度单色": ['#000000', '#333333', '#666666', '#999999', '#CCCCCC', '#E5E5E5'],
    "🌈 彩虹渐变": "husl",
    "🧬 科研柔和 (Sci)": ['#B4C7E7', '#E6B8D1', '#C5E0B4', '#FFE699', '#F4B183', '#D9D9D9'],
    "📑 NCD/HFD (Grey/Red)": ['#808080', '#D62728'],
    "✏️ 自定义...": None  # Triggers custom input
}

CUSTOM_PALETTE_KEY = "✏️ 自定义..."

def get_palette_colors(name, n_colors=None, custom_colors=None):
    """Get palette colors by name. If custom, parse from custom_colors string."""
    if name == CUSTOM_PALETTE_KEY and custom_colors:
        # Parse custom hex colors
        try:
            colors = [c.strip() for c in custom_colors.split(',') if c.strip().startswith('#')]
            if colors:
                return colors[:n_colors] if n_colors else colors
        except:
            pass
        return sns.color_palette("Set2", n_colors=n_colors)  # Fallback
    
    if name in PALETTES:
        param = PALETTES[name]
        if param is None:
            return sns.color_palette("Set2", n_colors=n_colors)
        elif isinstance(param, list):
            return param[:n_colors] if n_colors else param
        else:
            return sns.color_palette(param, n_colors=n_colors)
    return sns.color_palette("Set2", n_colors=n_colors)

# --- Setup Style ---
try:
    from tools.common import nature_style
except ImportError:
    class nature_style:
        @staticmethod
        def apply_nature_style():
            plt.rcParams['font.family'] = 'sans-serif'
            plt.rcParams['font.sans-serif'] = ['Arial', 'DejaVu Sans']
            sns.set_style("ticks")
            sns.set_context("paper")

st.set_page_config(page_title="ElementPrism - Visual Workbench", layout="wide", initial_sidebar_state="expanded")

# --- ElementPrism Brand Theme (Emerald Green) ---
st.markdown("""
<style>
    /* --- Global Theme --- */
    :root {
        --primary: #10b981;
        --bg-dark: #0d1117;
        --glass-bg: rgba(22, 27, 34, 0.7);
        --glass-border: 1px solid rgba(48, 54, 61, 0.7);
        --text-main: #e6edf3;
    }

    /* Force Dark Background */
    .stApp {
        background-color: var(--bg-dark);
        color: var(--text-main);
    }

    /* --- Inputs & Widgets (Glassmorphism) --- */
    /* Text Input, Selectbox, Number Input overlays */
    .stTextInput > div > div, 
    .stSelectbox > div > div, 
    .stNumberInput > div > div,
    .stMultiSelect > div > div {
        background-color: rgba(13, 17, 23, 0.5) !important;
        border: 1px solid rgba(16, 185, 129, 0.2) !important;
        border-radius: 8px !important;
        color: #fff !important;
    }
    
    /* Focus State */
    .stTextInput > div > div:focus-within, 
    .stSelectbox > div > div:focus-within {
        border-color: #10b981 !important;
        box_shadow: 0 0 0 1px #10b981 !important;
    }

    /* Sliders */
    .stSlider [data-baseweb="slider"] {
        margin-top: 1rem;
    }
    
    /* Buttons (Emerald Glow) */
    .stButton > button {
        background: rgba(16, 185, 129, 0.1) !important;
        border: 1px solid rgba(16, 185, 129, 0.5) !important;
        color: #10b981 !important;
        border-radius: 8px !important;
        transition: all 0.3s ease;
    }
    .stButton > button:hover {
        background: #10b981 !important;
        color: #fff !important;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.3);
    }
    
    /* Primary Action Buttons */
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
        color: white !important;
        border: none !important;
        font-weight: bold;
    }

    /* --- Containers & Expanders --- */
    /* Make Expanders look like Cards */
    .streamlit-expanderHeader {
        background-color: rgba(22, 27, 34, 0.6) !important;
        border: var(--glass-border) !important;
        border-radius: 8px !important;
        color: var(--text-main) !important;
    }
    .streamlit-expanderContent {
        background-color: rgba(22, 27, 34, 0.3) !important;
        border-left: 1px solid rgba(16, 185, 129, 0.2);
        border-right: 1px solid rgba(16, 185, 129, 0.2);
        border-bottom: 1px solid rgba(16, 185, 129, 0.2);
        border-bottom-left-radius: 8px;
        border-bottom-right-radius: 8px;
        color: var(--text-main) !important;
    }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #161b22;
        border-right: 1px solid rgba(48, 54, 61, 0.5);
    }
    
    /* Headings */
    h1, h2, h3 {
        color: #e6edf3 !important;
        font-family: 'Segoe UI', Inter, sans-serif !important;
    }
    h2 {
        border-bottom: 2px solid rgba(16, 185, 129, 0.3);
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem;
    }

    /* Tables/Dataframes */
    [data-testid="stDataFrame"] {
        border: 1px solid rgba(48, 54, 61, 0.5);
        border-radius: 8px;
    }

</style>
""", unsafe_allow_html=True)

# --- ElementPrism Header ---
# Using columns to create a balanced header
h_c1, h_c2 = st.columns([3, 1])
with h_c1:
    st.title("ElementPrism :rainbow[Workbench]")
    st.caption("One-Click Publication-Ready Figures for Biologists")


# --- Session State for Report ---
if 'report_items' not in st.session_state:
    st.session_state['report_items'] = []

def add_to_report(title, fig=None, df=None, text=None):
    """Add item to analysis report"""
    # Convert Plot to Image Bytes immediately to avoid state issues
    img_bytes = None
    if fig:
        buf = io.BytesIO()
        try:
            if hasattr(fig, 'savefig'):
                fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
            else:
                fig.figure.savefig(buf, format='png', dpi=300, bbox_inches='tight')
            img_bytes = buf.getvalue()
        except Exception as e:
            st.error(f"Error saving figure for report: {e}")

    item = {
        "title": title,
        "img_bytes": img_bytes,
        "df": df,
        "text": text,
        "time": datetime.datetime.now().strftime("%H:%M:%S")
    }
    st.session_state['report_items'].append(item)
    st.toast(f"✅ 已添加到报告: {title}")

# --- Helper Functions ---
def load_data(uploaded_file):
    try:
        if uploaded_file.name.endswith('.csv'):
            return pd.read_csv(uploaded_file)
        else:
            return pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"文件读取失败: {e}")
        return None

def get_download_buttons(fig, prefix, key_suffix, df_stats=None, report_title=None):
    c1, c2, c3 = st.columns([3, 1, 1])
    with c1:
        custom_name = st.text_input("文件名", value=f"{prefix}_{datetime.datetime.now().strftime('%Y%m%d')}", key=f"fname_{key_suffix}")
    
    # Report Button
    with c3:
        st.write("") # Spacer
        st.write("")
        if st.button("➕ 添加到报告", key=f"rpt_{key_suffix}"):
            add_to_report(report_title or prefix, fig=fig, df=df_stats)
    
    # Save Logic
    buf_pdf = io.BytesIO()
    buf_png = io.BytesIO()
    
    try:
        if hasattr(fig, 'savefig'):
            fig.savefig(buf_pdf, format="pdf", bbox_inches='tight')
            fig.savefig(buf_png, format="png", dpi=300, bbox_inches='tight')
        else:
            fig.figure.savefig(buf_pdf, format="pdf", bbox_inches='tight')
            fig.figure.savefig(buf_png, format="png", dpi=300, bbox_inches='tight')
    except: pass
    
    c_d1, c_d2 = st.columns(2)
    c_d1.download_button("📥 PDF", data=buf_pdf, file_name=f"{custom_name}.pdf", mime="application/pdf", key=f"dl_pdf_{key_suffix}")
    c_d2.download_button("📥 PNG", data=buf_png, file_name=f"{custom_name}.png", mime="image/png", key=f"dl_png_{key_suffix}")

# --- Educational Modules ---
def render_stat_guide():
    with st.expander("📚 ElementPrism导览：我该选什么检验？(Statistical Guide)", expanded=False):
        st.markdown("""
        ### Q1: 我有几组数据？
        
        #### 👉 只有 2 组 (例如: Control vs Treat)
        *   **Student's t-test (t检验)**: 
            *   数据符合**正态分布** (钟形曲线)。
            *   两组数据的**方差相似** (胖瘦差不多)。
        *   **Welch's t-test (校正t检验)**: 
            *   数据符合正态分布。
            *   但是**方差不齐** (一组很散，一组很聚)。*本工具会自动检测并选用此项。*
        *   **Mann-Whitney U test (非参数检验)**: 
            *   数据**不符合**正态分布 (分布很怪，或者样本量极少 <3)。
        
        #### 👉 有 3 组或更多 (例如: Control, Dose1, Dose2)
        *   **One-way ANOVA (单因素方差分析)**: 
            *   数据符合**正态分布**。
            *   **方差齐**。
            *   *后续会配合 Tukey HSD 进行两两比较。*
        *   **Kruskal-Wallis test (非参数方差分析)**: 
            *   数据**不符合**正态分布。
            *   *后续会配合 Dunn's test 进行两两比较。*
            
        ---
        ### 💡 常见名词解释
        *   **正态性 (Normality)**: 数据是否呈现中间高、两边低的对称分布。
        *   **方差齐性 (Homogeneity of Variance)**: 不同组别的数据离散程度是否一致。
        *   **P < 0.05**: 通常认为具有"统计学显著差异" (Statistically Significant)。
        *   **ns**: Not Significant，无显著差异 (P > 0.05)。
        """)

# --- Modules ---

def render_difference_module(data):
    st.header("📊 箱线图 (Boxplot)")
    render_stat_guide() # Insert Guide Here
    
    cols = data.columns.tolist()
    c1, c2, c3 = st.columns(3)
    x_col = c1.selectbox("分组列 (Group)", cols, index=0, key="diff_x")
    y_col = c2.selectbox("数值列 (Value)", cols, index=min(1, len(cols)-1), key="diff_y")
    hue_col = c3.selectbox("颜色分组 (Hue, 可选)", ["无"] + cols, index=0, key="diff_hue")

    group_order = st.multiselect("分组顺序", sorted(data[x_col].unique()), default=sorted(data[x_col].unique()))
    if not group_order: return

    # User input for plot title
    c_opt1, c_opt2, c_opt3 = st.columns([1, 1, 1])
    # Advanced Layout Option
    metric_as_x = c_opt1.checkbox("将数值列名作为X轴 (Single Metric)", value=False, key="diff_metric_x", help="开启后，X轴显示参数名，分组显示在图例中")
    # Palette Selection
    # Try to set default to Scientific if available
    palette_keys = list(PALETTES.keys())
    default_idx = next((i for i, k in enumerate(palette_keys) if "🧬" in k), 0)
    palette_name = c_opt3.selectbox("配色方案", palette_keys, index=default_idx, key="diff_palette")
    
    # Advanced Styling
    # Advanced Styling (Direct layout to fix Stlite issues)
    st.markdown("#### 🎨 样式与标注设置")
    c_s1, c_s2, c_s3, c_s4 = st.columns(4)
    show_points = c_s1.checkbox("显示散点 (Points)", value=True, key="diff_points")
    show_ns = c_s2.checkbox("显示 'ns'", value=False, key="diff_show_ns", help="勾选后将显示非显著差异 (p>0.05) 的标记")
    p_val_fmt = c_s3.selectbox("P值格式", ["Star (*)", "Simple (p=0.05)"], index=1, key="diff_pfmt")
    italic_xaxis = c_s4.checkbox("斜体X轴 (Italic)", value=True, key="diff_italic")
    
    # Width Input (Text fallback)
    default_width = "0.4" if metric_as_x else "0.5"
    bw_str = st.text_input("箱体宽度 (Width, 0.1-1.0)", default_width, key="diff_width")
    try: box_width = float(bw_str)
    except: box_width = float(default_width)
    
    # Custom color input (appears if user selects custom)
    custom_colors = None
    if palette_name == CUSTOM_PALETTE_KEY:
        custom_colors = st.text_input("输入自定义颜色 (逗号分隔)", placeholder="#FF0000, #00FF00, #0000FF", key="diff_custom_colors")

    custom_title = st.text_input("图表标题 (Plot Title)", value="", placeholder="留空则无标题", key="diff_plot_title")

    if st.button("🚀 运行分析 (智能模式)", type="primary"):
        st.session_state['diff_active'] = True
        
    if st.session_state.get('diff_active', False):
        st.divider()
        data_filtered = data[data[x_col].isin(group_order)].copy()
        
        # --- FIX: Type Consistency for Plotting & Stats ---
        data_filtered[x_col] = data_filtered[x_col].astype(str)
        group_order_str = [str(g) for g in group_order]
        
        # Robust Data Cleaning:
        data_filtered[y_col] = pd.to_numeric(data_filtered[y_col], errors='coerce')

        import numpy as np
        data_filtered = data_filtered.replace([np.inf, -np.inf], np.nan)
        data_filtered = data_filtered.dropna(subset=[y_col, x_col])

        # --- Layout Logic Transformation ---
        plot_x_col = x_col
        plot_hue_col = hue_col if hue_col != "无" else None
        plot_order = group_order_str
        hue_order = None
        
        if metric_as_x:
            # Single Metric Mode: X-axis = Metric Name, Hue = Original Group
            metric_name = y_col
            data_filtered["_Metric"] = metric_name
            plot_x_col = "_Metric"
            plot_order = [metric_name]
            
            plot_hue_col = x_col # Original group becomes Hue
            hue_order = group_order_str # Original group order becomes Hue order
        elif plot_hue_col:
             hue_order = sorted(data_filtered[plot_hue_col].astype(str).unique())

        # Prepare Data Groups for Stats (Based on original logic or new logic?)
        # Original logic was comparing X groups.
        # If Metric Mode, we want to compare Hue groups (which was original X).
        # So we can keep using 'groups_data' based on original X_col for global stats.
        groups_data = [data_filtered[data_filtered[x_col]==g][y_col] for g in group_order_str]
        
        if len(groups_data) < 2:
            st.error("需要至少两组数据进行比较")
            return

        # --- 1. Assumption Checks ---
        st.subheader("🧐 统计策略检测 (Statistical Strategy)")
        
        normality_passed = True
        dataset_too_small = False
        shapiro_res = {}
        
        # Normality (Shapiro-Wilk)
        for g, vals in zip(group_order_str, groups_data):
            if len(vals) < 3:
                dataset_too_small = True
                normality_passed = False 
                shapiro_res[g] = (0, 0, "N<3, Fail")
            else:
                stat_val, p = stats.shapiro(vals)
                shapiro_res[g] = (stat_val, p, "Pass" if p > 0.05 else "Fail")
                if p <= 0.05: normality_passed = False

        # Homogeneity (Levene)
        # Only check Levene if we have enough samples to potentially run ANOVA
        if len(groups_data) > 1 and all(len(g) >= 2 for g in groups_data):
            levene_stat, levene_p = stats.levene(*groups_data)
            equal_var = levene_p > 0.05
        else:
            levene_p = 0
            equal_var = False # Conservative fallback

        # Display Report
        with st.expander("📄 查看详细假设检验报告 (Assumption Check Report)", expanded=True):
            c_r1, c_r2 = st.columns(2)
            with c_r1:
                st.markdown("**1. 正态性检验 (Shapiro-Wilk)**")
                for g, res in shapiro_res.items():
                    color = "green" if res[2] == "Pass" else "red"
                    val_str = f"p={res[1]:.4f}" if isinstance(res[1], float) else res[2]
                    st.markdown(f"- {g}: :{color}[{val_str}]")
                if dataset_too_small: st.warning("⚠️ 样本量 < 3，自动视为非正态")
            
            with c_r2:
                st.markdown("**2. 方差齐性检验 (Levene)**")
                color = "green" if equal_var else "red"
                st.markdown(f"- P-Value: :{color}[{levene_p:.4f}]")
                st.markdown(f"- 结论: {'方差齐 (Equal Var)' if equal_var else '方差不齐 (Unequal Var)'}")

        # --- 2. Select Test ---
        test_name = "Unknown"
        p_global = 1.0
        statistic = 0.0
        sig_pairs = []

        if len(groups_data) == 2:
            g1, g2 = groups_data[0], groups_data[1]
            if normality_passed:
                if equal_var:
                    test_name = "Student's t-test"
                    statistic, p_global = stats.ttest_ind(g1, g2, equal_var=True)
                else:
                    test_name = "Welch's t-test"
                    statistic, p_global = stats.ttest_ind(g1, g2, equal_var=False)
            else:
                test_name = "Mann-Whitney U"
                statistic, p_global = stats.mannwhitneyu(g1, g2)
            
            if p_global < 0.05 or show_ns:
                sig_pairs = [((group_order_str[0], group_order_str[1]), p_global)]

        else:
            if normality_passed and equal_var:
                test_name = "One-way ANOVA"
                statistic, p_global = stats.f_oneway(*groups_data)
                st.info(f"👉 数据符合正态分布且方差齐，自动选择: **{test_name} (Post-hoc: Tukey HSD)**")
                
                # Checkbox to force running post-hoc even if global p > 0.05 is generally discouraged, 
                # but if user wants to see "ns", we should allow pairwise checks.
                # Standard practice: Only run post-hoc if global p < 0.05.
                # If show_ns is True, user implies they want to see the comparisons anyway.
                if p_global < 0.05 or show_ns:
                    try:
                        tukey = pairwise_tukeyhsd(endog=data_filtered[y_col], groups=data_filtered[x_col], alpha=0.05)
                        tukey_data = pd.DataFrame(data=tukey.summary().data[1:], columns=tukey.summary().data[0])
                        for _, row in tukey_data.iterrows():
                             if row['reject'] or show_ns:
                                 sig_pairs.append( ((str(row['group1']), str(row['group2'])), row['p-adj']) )
                    except Exception as e:
                        st.warning(f"Tukey Post-hoc failed ({str(e)}), switching to Robust Pairwise T-tests (Holm-Sidak)...")
                        # Fallback: Pairwise T-test with Holm correction
                        from statsmodels.stats.multitest import multipletests
                        p_vals = []
                        pair_indices = []
                        
                        # Calculate raw p-values
                        for i in range(len(group_order_str)):
                            for j in range(i+1, len(group_order_str)):
                                g1 = groups_data[i]
                                g2 = groups_data[j]
                                # Use Welch's T-test for robustness in fallback
                                _, p = stats.ttest_ind(g1, g2, equal_var=False) 
                                p_vals.append(p)
                                pair_indices.append((group_order_str[i], group_order_str[j]))
                        
                        # Apply Correction (Holm)
                        if p_vals:
                            reject, p_adjusted, _, _ = multipletests(p_vals, method='holm')
                            for k, (pai, padj) in enumerate(zip(pair_indices, p_adjusted)):
                                if reject[k] or show_ns:
                                    sig_pairs.append((pai, padj))
            else:
                test_name = "Kruskal-Wallis"
                statistic, p_global = stats.kruskal(*groups_data)
                st.info(f"👉 数据不符合正态分布或方差不齐，自动选择: **{test_name} (Post-hoc: Dunn)**")
                
                if p_global < 0.05 or show_ns:
                    dunn = posthoc_dunn(data_filtered, val_col=y_col, group_col=x_col, p_adjust='holm')
                    for i in range(len(group_order_str)):
                        for j in range(i+1, len(group_order_str)):
                            p_val = dunn.loc[group_order_str[i], group_order_str[j]]
                            if p_val < 0.05 or show_ns:
                                sig_pairs.append( ((group_order_str[i], group_order_str[j]), p_val) )

        if len(groups_data) == 2:
             st.info(f"👉 自动选择: **{test_name}**")

        # --- 3. Plotting ---
        nature_style.apply_nature_style()
        fig, ax = plt.subplots(figsize=(5, 5))
        
        # Get palette colors
        # If Hue is used, we need colors for Hue levels.
        # If Metric Mode, Hue=Original X, so n_colors=len(group_order)
        # If Standard Mode + Hue, n_colors=len(hue_unique)
        # If Standard Mode No Hue, n_colors=len(group_order)
        if plot_hue_col:
            hue_levels = hue_order if hue_order else sorted(data_filtered[plot_hue_col].astype(str).unique())
            n_colors = len(hue_levels)
        else:
            n_colors = len(group_order_str)
            
        colors = get_palette_colors(palette_name, n_colors=n_colors, custom_colors=custom_colors)
        
        if metric_as_x:
            real_width = 0.8 # Standard width for dodged items
            padding_factor = (1.1 - box_width) * 1.5 
            ax.set_xlim(-0.5 - padding_factor, 0.5 + padding_factor)
            # Metric mode implies dodging effectively since X is the metric
            should_dodge = True 
        else:
            real_width = box_width
            # Dodge only if hue is present (and not just coloring by X)
            should_dodge = True if plot_hue_col else False

        # Boxplot
        sns.boxplot(x=plot_x_col, y=y_col, hue=plot_hue_col, data=data_filtered, 
                    order=plot_order, hue_order=hue_order,
                    width=real_width, ax=ax, palette=colors,
                    dodge=should_dodge, # Explicit dodge control
                    linewidth=1.0, fliersize=0) 
        
        # Stripplot (Styled)
        if show_points:
            sns.stripplot(x=plot_x_col, y=y_col, hue=plot_hue_col, data=data_filtered, 
                          order=plot_order, hue_order=hue_order,
                          dodge=should_dodge, # Match boxplot dodge setting exactly
                          size=5, jitter=True, ax=ax,
                          color="white", edgecolor="gray", linewidth=1,
                          legend=False)
        
        # Annotate
        text_format = 'simple' if 'Simple' in p_val_fmt else 'star'
        
        # Stats Annotation Logic Update for Layouts
        try:
            # Case 1: Standard Mode (No Hue) -> Compare X groups
            if not plot_hue_col:
                if sig_pairs:
                     plot_pairs = [p[0] for p in sig_pairs]
                     p_values = [p[1] for p in sig_pairs]
                     annotator = Annotator(ax, plot_pairs, data=data_filtered, x=plot_x_col, y=y_col, order=plot_order)
                     annotator.configure(test=None, text_format=text_format, loc='inside' if text_format=='star' else 'outside', verbose=False)
                     annotator.set_pvalues(p_values)
                     annotator.annotate()
            
            # Case 2: Metric as X Mode -> Compare Hue groups within the single Metric X
            elif metric_as_x:
                # We want to compare the Hue groups (which are the original X groups)
                # sig_pairs contains (GroupA, GroupB)
                # We need to construct pairs like: (("Metric", GroupA), ("Metric", GroupB))
                if sig_pairs:
                     # Reformat pairs for statannotations with hue
                     hue_plot_pairs = []
                     for (g1, g2), _ in sig_pairs:
                         hue_plot_pairs.append(((metric_name, g1), (metric_name, g2)))
                     
                     p_values = [p[1] for p in sig_pairs]
                     
                     annotator = Annotator(ax, hue_plot_pairs, data=data_filtered, x=plot_x_col, y=y_col, hue=plot_hue_col, 
                                           order=plot_order, hue_order=hue_order)
                     annotator.configure(test=None, text_format=text_format, loc='inside' if text_format=='star' else 'outside', verbose=False)
                     annotator.set_pvalues(p_values)
                     annotator.annotate()
            
            # Case 3: Standard Mode + Hue -> Hue comparisons? Or X comparisons?
            # Typically user wants to compare Hue within X. Not implemented in previous stats logic but useful.
            # For now, let's skip complex stats for Case 3 to avoid breakage, unless user asks.
            
        except Exception as e:
             st.write(f"Annotation Warning: {e}")
        
        if custom_title:
            ax.set_title(custom_title, fontsize=12)
        else:
            ax.set_title("") 

        if italic_xaxis and not metric_as_x: # If metric as X, usually not italic
             ax.set_xticklabels(ax.get_xticklabels(), fontstyle='italic')

        sns.despine()
        
        # Move Legend
        if plot_hue_col:
             plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)

        st.pyplot(fig)
        
        # Result DataFrame
        res_df = pd.DataFrame({
            "Test Selected": [test_name], 
            "Statistic": [statistic], 
            "Global P-Value": [p_global],
            "Normality": ["Pass" if normality_passed else "Fail"],
            "Equal Variance": ["Pass" if equal_var else "Fail"]
        })
        st.dataframe(res_df)
        
        get_download_buttons(fig, f"Stat_{y_col}", "stat", df_stats=res_df, report_title=f"Analysis: {test_name} on {y_col}")

# --- Bar Chart Module ---
def render_barplot_module(data):
    st.header("📊 条形图 (Bar Chart)")
    st.markdown("绘制带误差棒的条形图，支持统计分析和显著性标注。")
    
    cols = data.columns.tolist()
    num_cols = data.select_dtypes(include=['number']).columns.tolist()
    
    c1, c2, c3 = st.columns(3)
    c1, c2, c3 = st.columns(3)
    x_col = c1.selectbox("分组列 (X)", cols, index=0, key="bar_x")
    # Change to multiselect to support Wide Mode
    y_cols = c2.multiselect("数值列 (Y, 可多选)", num_cols, default=[num_cols[0]] if num_cols else None, key="bar_y")
    hue_col = c3.selectbox("颜色分组 (Hue, 可选)", ["无"] + cols, index=0, key="bar_hue")
    
    if not y_cols:
        st.warning("请至少选择一个数值列")
        return
    
    # Options Row 1
    c_opt1, c_opt2, c_opt3 = st.columns(3)
    agg_method = c_opt1.selectbox("聚合方式", ["mean", "median", "sum", "count"], key="bar_agg")
    error_type = c_opt2.selectbox("误差棒", ["sd", "se", "ci", "无"], key="bar_err")
    palette_keys = list(PALETTES.keys())
    default_idx = next((i for i, k in enumerate(palette_keys) if "🧬" in k), 0)
    palette_name = c_opt3.selectbox("配色方案", palette_keys, index=default_idx, key="bar_palette") 

    # Advanced Styling
    st.markdown("#### 🎨 样式与标注设置")
    c_s1, c_s2, c_s3, c_s4, c_s5 = st.columns(5)
    show_points = c_s1.checkbox("显示散点", value=True, key="bar_points")
    show_edge_color = c_s2.checkbox("显示黑边框", value=True, key="bar_edge")
    p_val_fmt = c_s3.selectbox("P值格式", ["Star (*)", "Simple (p=0.05)"], index=0, key="bar_pfmt")
    italic_xaxis = c_s4.checkbox("斜体X轴", value=True, key="bar_italic")
    show_ns = c_s5.checkbox("显示 'ns'", value=False, key="bar_show_ns")
    
    # Extra Aesthetics Row
    c_e1, c_e2, c_e3 = st.columns(3)
    y_fs_str = c_e1.text_input("Y轴字体大小 (Font Size)", "12", key="bar_yfs")
    sig_lw_str = c_e2.text_input("显著性线宽 (Line Width)", "1.5", key="bar_siglw")
    lh_str = c_e3.text_input("显著性线高 (Line Height, 0=直线)", "0.02", key="bar_lh")
    try: y_fontsize = float(y_fs_str)
    except: y_fontsize = 12
    try: sig_linewidth = float(sig_lw_str)
    except: sig_linewidth = 1.5
    try: line_height_val = float(lh_str)
    except: line_height_val = 0.02
    
    use_shapes = st.checkbox("散点使用不同形状 (Group Shapes)", value=False, key="bar_shapes")
    
    # Layout Option
    # If multiple Y selected, force metric_as_x behavior internally
    is_wide_mode = len(y_cols) > 1
    if is_wide_mode:
        st.info("💡 检测到多指标模式：已自动将指标名称设为 X 轴，原 '分组列 (X)' 已被忽略，请使用 '颜色分组' 进行群组区分。")
    
    metric_as_x = st.checkbox("将数值列名作为X轴 (Single Metric)", value=False, key="bar_metric_x", disabled=is_wide_mode, help="单指标时可选。开启后，X轴显示参数名。")
    
    # Width Input
    default_bar_width = "0.4" if metric_as_x else "0.6"
    bw_str = st.text_input("条形宽度 (Width, 0.1-1.0)", default_bar_width, key="bar_width")
    try: bar_width_val = float(bw_str)
    except: bar_width_val = float(default_bar_width)

    # Custom color input
    custom_colors = None
    if palette_name == CUSTOM_PALETTE_KEY:
        custom_colors = st.text_input("输入自定义颜色 (逗号分隔)", placeholder="#FF0000, #00FF00, #0000FF", key="bar_custom_colors")
    
    custom_title = st.text_input("图表标题", value="", placeholder="留空则无标题", key="bar_plot_title")
    custom_ylabel = st.text_input("Y轴标题 (自定义)", value="", placeholder="例如: Weight (g) 或 Expression Level", key="bar_ylabel")
    
    # Group order
    group_order = st.multiselect("分组顺序 (X)", sorted(data[x_col].unique()), default=sorted(data[x_col].unique()), key="bar_order")
    if not group_order: return
    
    if st.button("🚀 生成条形图", type="primary", key="bar_btn"):
        st.session_state['bar_active'] = True
    
    show_debug = st.checkbox("显示调试信息 (Debug Info)", value=False, key="bar_debug")

    if st.session_state.get('bar_active', False):
        st.divider()
        import numpy as np
        
        # Data Preparation
        if is_wide_mode:
            # Wide Mode: Melt Data
            # ID vars = hue_col if exists
            id_vars = [hue_col] if hue_col != "无" else []
            # We must keep row index to align valid data if needed, but melt usually handles it.
            # Filter down to useful cols first
            subset = data[id_vars + y_cols].copy()
            data_filtered = subset.melt(id_vars=id_vars, value_vars=y_cols, var_name="_Metric", value_name="_Value")

            # Setup Plot Cols
            plot_x_col = "_Metric"
            y_col = "_Value" # Override y_col for plotting
            plot_x_col = "_Metric"
            y_col = "_Value" # Override y_col for plotting
            plot_hue_col = hue_col if hue_col != "无" else None
            
            # Key Fix: Ensure Hue Column is String
            if plot_hue_col:
                data_filtered[plot_hue_col] = data_filtered[plot_hue_col].astype(str)
            
            plot_order = y_cols # The order of X is the selection order
            hue_order = sorted(data_filtered[plot_hue_col].astype(str).unique()) if plot_hue_col else None
            
            metric_as_x = True # Force this flag for logic downstream
            
        else:
            # --- Standard Mode Logic ---
            y_col = y_cols[0] # Define y_col from the list for standard mode
            data_filtered = data[data[x_col].isin(group_order)].copy()
            data_filtered[x_col] = data_filtered[x_col].astype(str)
            group_order_str = [str(g) for g in group_order]
            
            plot_x_col = x_col
            plot_hue_col = hue_col if hue_col != "无" else None
            plot_order = group_order_str
            
            # Sub-branches for Standard Mode layouts
            if metric_as_x:
                # Single Metric Mode (Standard Data Structure)
                metric_name = y_col
                data_filtered["_Metric"] = metric_name
                plot_x_col = "_Metric"
                plot_order = [metric_name]
                
                plot_hue_col = x_col # Original Group becomes Hue
                hue_order = group_order_str
            else:
                # True Standard Barplot
                hue_order = sorted(data_filtered[plot_hue_col].astype(str).unique()) if plot_hue_col else None


        
        # Aggregation
        if agg_method == "mean":
            estimator = np.mean
        elif agg_method == "median":
            estimator = np.median
        elif agg_method == "sum":
            estimator = np.sum
        else:
            estimator = len
        
        # Error Bar
        errorbar_param = None if error_type == "无" else error_type
        
        # Plotting Setup
        nature_style.apply_nature_style()
        fig, ax = plt.subplots(figsize=(5, 5))
        
        hue = None if hue_col == "无" else hue_col
        # Calculate N colors needed
        if plot_hue_col:
            hue_levels = hue_order if hue_order else sorted(data_filtered[plot_hue_col].astype(str).unique())
            n_colors = len(hue_levels)
        else:
            n_colors = len(plot_order)
        
        # Initialize analysis log
        analysis_log = []
        analysis_log.append(f"**分析模式**: {'多指标 (Wide Mode)' if is_wide_mode else '标准 (Standard)'}")
        analysis_log.append(f"**X轴**: {plot_x_col}")
        if plot_hue_col: analysis_log.append(f"**分组变量 (Hue)**: {plot_hue_col}")

        colors = get_palette_colors(palette_name, n_colors=n_colors, custom_colors=custom_colors)
        
        if metric_as_x:
            real_width = 0.8 
            # Only restrict X-limit if we have a SINGLE metric, to keep it centered and not too wide.
            # If we have multiple metrics, let matplotlib auto-scale.
            if len(plot_order) == 1:
                padding_factor = (1.1 - bar_width_val) * 1.5
                ax.set_xlim(-0.5 - padding_factor, 0.5 + padding_factor)
            should_dodge = True
        else:
            # Standard Width Control
            real_width = bar_width_val
            # Dodge only if actual Hue group exists
            should_dodge = True if plot_hue_col else False
            
            # Special Case: Seaborn barplot defaults to dodge=True even without hue
            # We must force it off if we want centered bars
            if not plot_hue_col: should_dodge = False

        # 1. Main Bar Plot
        edge_col = "black" if show_edge_color else None
        lw = 1.0 if show_edge_color else 0
        sns.barplot(
            data=data_filtered, x=plot_x_col, y=y_col, order=plot_order,
            hue=plot_hue_col, hue_order=hue_order,
            estimator=estimator, errorbar=errorbar_param,
            palette=colors, ax=ax, capsize=0.1, errwidth=1.5,
            edgecolor=edge_col, linewidth=lw, 
            width=real_width, dodge=should_dodge # Unified Dodge
        )
        
        # 2. Points Overlay
        if show_points:
            try:
                if use_shapes and plot_hue_col and hue_order:
                     marker_list = ['o', 's', '^', 'D', 'v', '<', '>', 'p', '*', 'h', 'H', '+', 'x', 'X', 'd', '|', '_']
                     
                     for i, h_val in enumerate(hue_order):
                         m = marker_list[i % len(marker_list)]
                         subset_for_hue = data_filtered[data_filtered[plot_hue_col] == h_val]
                         
                         sns.stripplot(
                            data=subset_for_hue, x=plot_x_col, y=y_col, 
                            hue=plot_hue_col, hue_order=hue_order, order=plot_order,
                            dodge=should_dodge, ax=ax, palette=colors,
                            jitter=True, size=5, linewidth=1, edgecolor='gray', marker=m, alpha=0.7, legend=False
                         )
                else:
                     sns.stripplot(
                        data=data_filtered, x=plot_x_col, y=y_col, 
                        hue=plot_hue_col, hue_order=hue_order, order=plot_order,
                        dodge=should_dodge, ax=ax, palette=colors,
                        jitter=True, size=5, linewidth=1, edgecolor='gray', marker='o', alpha=0.7, legend=False
                     )
            except Exception as e:
                # Fallback
                if show_debug: st.error(f"Shape Error: {e}")
                sns.stripplot(
                    data=data_filtered, x=plot_x_col, y=y_col, 
                    hue=plot_hue_col, hue_order=hue_order, order=plot_order,
                    dodge=should_dodge, ax=ax, palette=colors,
                    jitter=True, size=5, linewidth=1, edgecolor='gray', marker='o', alpha=0.7, legend=False
                )

        # 3. Statistical Analysis
        text_format = 'simple' if 'Simple' in p_val_fmt else 'star'
        
        # --- Logic for Stats ---
        try:
            if is_wide_mode:
                # Wide Mode Stats: Compare Hue groups within each Metric (X)
                if plot_hue_col and hue_order and len(hue_order) >= 2:
                    analysis_log.append(f"\n**[统计检验] 多指标模式 (Wide Mode)**")
                    hue_plot_pairs = []
                    p_values = []
                    import itertools
                    
                    # Iterate over each Metric on X-axis
                    for metric in plot_order: # plot_order is list of metrics
                        # Filter data for this metric
                        metric_data = data_filtered[data_filtered[plot_x_col] == metric]
                        
                        # Check number of groups
                        valid_hue_data = []
                        valid_hue_names = []
                        for h in hue_order:
                            v = metric_data[metric_data[plot_hue_col] == h][y_col].dropna()
                            if len(v) > 1 and v.nunique() > 0:
                                valid_hue_data.append(v)
                                valid_hue_names.append(h)
                        
                        if len(valid_hue_data) < 2:
                             if show_debug: st.text(f"Skip {metric}: Not enough valid groups ({len(valid_hue_data)})")
                             continue
                             
                        # Branch: 2 Groups (MWU) vs >2 Groups (Kruskal+Dunn)
                        if len(valid_hue_data) == 2:
                             # MWU
                             analysis_log.append(f"- *{metric}*: 检测到2组 ({valid_hue_names}), 使用 **Mann-Whitney U 检验**")
                             h1, h2 = valid_hue_names[0], valid_hue_names[1]
                             v1, v2 = valid_hue_data[0], valid_hue_data[1]
                             try:
                                 if v1.nunique() == 1 and v2.nunique() == 1 and v1.iloc[0] == v2.iloc[0]:
                                     continue
                                 _, p = stats.mannwhitneyu(v1, v2)
                                 if show_debug: st.text(f"MWU {metric} {h1}v{h2}: p={p:.4f}")
                                 if p < 0.05 or show_ns:
                                     hue_plot_pairs.append(((metric, h1), (metric, h2)))
                                     p_values.append(p)
                             except Exception as e:
                                 if show_debug: st.error(f"MWU Error {metric}: {e}")

                        else:
                             # Kruskal + Dunn
                             analysis_log.append(f"- *{metric}*: 检测到{len(valid_hue_data)}组 ({valid_hue_names}), 使用 **Kruskal-Wallis + Dunn's Post-hoc**")
                             try:
                                 _, k_p = stats.kruskal(*valid_hue_data)
                                 if show_debug: st.text(f"Kruskal {metric}: p={k_p:.4f}")
                                 
                                 if k_p < 0.05 or show_ns:
                                     # Post-hoc Dunn
                                     # Need to reconstruct a subset dataframe for scikit-posthocs
                                     # Filter metric data to only valid hues
                                     sub_df = metric_data[metric_data[plot_hue_col].isin(valid_hue_names)].copy()
                                     dunn = posthoc_dunn(sub_df, val_col=y_col, group_col=plot_hue_col, p_adjust='holm')
                                     
                                     import itertools
                                     pairs = list(itertools.combinations(valid_hue_names, 2))
                                     
                                     for h1, h2 in pairs:
                                         try:
                                             p_val = dunn.loc[h1, h2]
                                             if show_debug: st.text(f"Dunn {metric} {h1}v{h2}: p={p_val:.4f}")
                                             if p_val < 0.05 or show_ns:
                                                 hue_plot_pairs.append(((metric, h1), (metric, h2)))
                                                 p_values.append(p_val)
                                         except: pass
                             except Exception as e:
                                 if show_debug: st.error(f"KW/Dunn Error {metric}: {e}")
                    
                    if hue_plot_pairs:
                        annotator = Annotator(ax, hue_plot_pairs, data=data_filtered, x=plot_x_col, y=y_col, hue=plot_hue_col, 
                                              order=plot_order, hue_order=hue_order)
                        annotator.configure(test=None, text_format=text_format, loc='inside' if text_format=='star' else 'outside', verbose=False, line_offset=0.08, line_offset_to_group=0.05, line_width=sig_linewidth, line_height=line_height_val)
                        annotator.set_pvalues(p_values)
                        annotator.annotate()

            elif not plot_hue_col:
                # Simple comparisons between X groups
                # Ensure we drop NaNs
                groups_data = [data_filtered[data_filtered[x_col]==g][y_col].dropna() for g in group_order_str]
                
                # Filter out empty or too small groups
                valid_groups_indices = [i for i, g in enumerate(groups_data) if len(g) > 1 and g.nunique() > 0]
                
                if len(valid_groups_indices) >= 2:
                    sig_pairs_found = []
                    # KW or Mann-Whitney
                    try:
                        if len(valid_groups_indices) == 2:
                            # Compare the two valid groups
                            idx1, idx2 = valid_groups_indices[0], valid_groups_indices[1]
                            g1_data, g2_data = groups_data[idx1], groups_data[idx2]
                            
                            if g1_data.nunique() == 1 and g2_data.nunique() == 1 and g1_data.iloc[0] == g2_data.iloc[0]:
                                pass # Constant identical
                            else:
                                _, p = stats.mannwhitneyu(g1_data, g2_data)
                                if p < 0.05 or show_ns:
                                    sig_pairs_found.append(((group_order_str[idx1], group_order_str[idx2]), p))
                        else:
                            # Kruskal for multiple groups
                            valid_data_list = [groups_data[i] for i in valid_groups_indices]
                            _, k_p = stats.kruskal(*valid_data_list)
                            
                            if k_p < 0.05 or show_ns:
                                dunn = posthoc_dunn(data_filtered, val_col=y_col, group_col=x_col, p_adjust='holm')
                                for i in range(len(group_order_str)):
                                    for j in range(i+1, len(group_order_str)):
                                        try:
                                            p_val = dunn.loc[group_order_str[i], group_order_str[j]]
                                            if p_val < 0.05 or show_ns:
                                                sig_pairs_found.append(((group_order_str[i], group_order_str[j]), p_val))
                                        except: pass
                    except Exception as e:
                         print(f"No Hue Stats error: {e}")
                         pass

                    
                    if sig_pairs_found:
                        annotator = Annotator(ax, [p[0] for p in sig_pairs_found], data=data_filtered, x=plot_x_col, y=y_col, order=plot_order)
                        annotator.configure(test=None, text_format=text_format, loc='inside' if text_format=='star' else 'outside', verbose=False, line_width=sig_linewidth, line_height=line_height_val)
                        annotator.set_pvalues([p[1] for p in sig_pairs_found])
                        annotator.annotate()

            elif metric_as_x:
                # Metric Mode: Compare Hue groups (Original Groups)
                # Reformat pairs for statannotations with hue
                # Compare all pairs in group_order
                sig_pairs_found = []
                groups_data = [data_filtered[data_filtered[x_col]==g][y_col] for g in group_order_str]
                
                # We need p-values for pairs
                # Prepare data list
                groups_data = []
                valid_groups = []
                hue_plot_pairs = []
                p_values = []
                
                for g in group_order_str:
                    v = data_filtered[data_filtered[x_col]==g][y_col].dropna()
                    if len(v) > 1 and v.nunique() > 0:
                        groups_data.append(v)
                        valid_groups.append(g)
                
                if len(groups_data) >= 2:
                    if len(groups_data) == 2:
                        # MWU
                        analysis_log.append(f"- *{metric_name}*: 检测到2组 ({valid_groups}), 使用 **Mann-Whitney U 检验**")
                        g1, g2 = valid_groups[0], valid_groups[1]
                        v1, v2 = groups_data[0], groups_data[1]
                        try:
                            if v1.nunique() == 1 and v2.nunique() == 1 and v1.iloc[0] == v2.iloc[0]:
                                pass
                            else:
                                _, p = stats.mannwhitneyu(v1, v2)
                                if show_debug: st.text(f"Metro MWU {g1}v{g2}: p={p:.4f}")
                                if p < 0.05 or show_ns:
                                    hue_plot_pairs.append(((metric_name, g1), (metric_name, g2)))
                                    p_values.append(p)
                        except Exception as e:
                            if show_debug: st.error(f"Metro MWU Error: {e}")
                    else:
                        # Kruskal
                        analysis_log.append(f"- *{metric_name}*: 检测到{len(groups_data)}组 ({valid_groups}), 使用 **Kruskal-Wallis + Dunn's Post-hoc**")
                        try:
                            _, k_p = stats.kruskal(*groups_data)
                            if show_debug: st.text(f"Metro Kruskal: p={k_p:.4f}")
                            if k_p < 0.05 or show_ns:
                                dunn = posthoc_dunn(data_filtered, val_col=y_col, group_col=x_col, p_adjust='holm')
                                import itertools
                                pairs = list(itertools.combinations(valid_groups, 2))
                                for g1, g2 in pairs:
                                    try:
                                        p_val = dunn.loc[g1, g2]
                                        if show_debug: st.text(f"Metro Dunn {g1}v{g2}: p={p_val:.4f}")
                                        if p_val < 0.05 or show_ns:
                                            hue_plot_pairs.append(((metric_name, g1), (metric_name, g2)))
                                            p_values.append(p_val)
                                    except: pass
                        except Exception as e:
                             if show_debug: st.error(f"Metro KW Error: {e}")

                if hue_plot_pairs:
                     annotator = Annotator(ax, hue_plot_pairs, data=data_filtered, x=plot_x_col, y=y_col, hue=plot_hue_col, 
                                           order=plot_order, hue_order=hue_order)
                     annotator.configure(test=None, text_format=text_format, loc='inside' if text_format=='star' else 'outside', verbose=False, line_offset=0.08, line_offset_to_group=0.05, line_width=sig_linewidth, line_height=line_height_val)
                     annotator.set_pvalues(p_values)
                     annotator.annotate()

            else:
                # Hue Stats (Standard)
                # Compare hues within each X group (Common requirement)
                hue_order = sorted(data_filtered[hue_col].unique())
                if len(hue_order) >= 2:
                    plot_pairs = []
                    p_values = []
                    
                    for g in group_order_str:
                         # Compare first two hues for each group
                         h1, h2 = hue_order[0], hue_order[1]
                         v1 = data_filtered[(data_filtered[x_col]==g) & (data_filtered[hue_col]==h1)][y_col].dropna()
                         v2 = data_filtered[(data_filtered[x_col]==g) & (data_filtered[hue_col]==h2)][y_col].dropna()
                         
                         if len(v1) > 1 and len(v2) > 1:
                              if v1.nunique() == 1 and v2.nunique() == 1 and v1.iloc[0] == v2.iloc[0]:
                                   continue
                              try:
                                  _, p = stats.mannwhitneyu(v1, v2)
                                  if p < 0.05 or show_ns:
                                      plot_pairs.append(((g, h1), (g, h2)))
                                      p_values.append(p)
                              except Exception as e:
                                  print(f"Standard Hue Stats error for {g}: {e}")
                                  pass
                    
                    if plot_pairs:
                        annotator = Annotator(ax, plot_pairs, data=data_filtered, x=x_col, y=y_col, hue=hue_col, order=group_order_str, hue_order=hue_order)
                        annotator.configure(test=None, text_format=text_format, loc='inside' if text_format=='star' else 'outside', verbose=False, line_width=sig_linewidth, line_height=line_height_val)
                        annotator.set_pvalues(p_values)
                        annotator.annotate()
        except Exception as e:
            st.warning(f"Stats Calculation Error: {e}")

        # Formatting
        if custom_title:
             ax.set_title(custom_title, fontsize=12)
        
        if italic_xaxis and not metric_as_x:
             ax.set_xticklabels(ax.get_xticklabels(), fontstyle='italic')
        
        # Rotate X labels if there are many
        if len(plot_order) > 4 or any(len(str(s)) > 10 for s in plot_order):
            plt.xticks(rotation=45, ha='right')
        
        # Axis Labels
        if custom_ylabel:
            ax.set_ylabel(custom_ylabel, fontsize=y_fontsize) # User override
        elif metric_as_x:
             ax.set_xlabel("") # Hide X Label since ticks verify it
             # If only 1 metric in wide/metric list, use its name. Else generic.
             if len(plot_order) == 1:
                 ax.set_ylabel(plot_order[0], fontsize=y_fontsize)
             else:
                 ax.set_ylabel("Value", fontsize=y_fontsize) 
        else:
             ax.set_xlabel(x_col, fontsize=y_fontsize)
             ax.set_ylabel(y_col, fontsize=y_fontsize)
        sns.despine()
        
        # Adjust Legend
        # If hue is active, place legend outside or top
        if plot_hue_col:
             plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)
        
        st.pyplot(fig)
        
        if show_debug and not analysis_log:
             st.info("Debugging active but no stats logs generated.")

        with st.expander("📊 统计分析方法报告 (Statistical Methodology)", expanded=True):
            for log in analysis_log:
                st.markdown(log)
            st.caption("注：非参数检验通常用于样本量较小或非正态分布的数据。P值校正采用Holm方法。")

        st.markdown("---")
        
        get_download_buttons(fig, f"Bar_{y_col}", "bar", report_title=f"Bar Chart: {y_col} by {x_col}")

# --- Line Chart Module ---
def render_linechart_module(data):
    st.header("📈 折线图 (Line Chart)")
    st.markdown("绘制折线图，支持多组数据趋势比较和误差带显示。")

    cols = data.columns.tolist()
    num_cols = data.select_dtypes(include=['number']).columns.tolist()

    c1, c2, c3 = st.columns(3)
    x_col = c1.selectbox("X轴 (Time/Category)", cols, index=0, key="line_x")
    y_col = c2.selectbox("Y轴 (Value)", num_cols, index=0, key="line_y")
    hue_col = c3.selectbox("分组 (Group, 可选)", ["无"] + cols, index=0, key="line_hue")

    # Options
    c_opt1, c_opt2, c_opt3 = st.columns(3)
    error_type = c_opt1.selectbox("误差带", ["sd", "se", "ci (95%)", "无"], index=0, key="line_err")
    show_points = c_opt2.checkbox("显示数据点", value=True, key="line_points")
    
    palette_keys = list(PALETTES.keys())
    default_idx = next((i for i, k in enumerate(palette_keys) if "Sci" in k), 0)
    palette_name = c_opt3.selectbox("配色方案", palette_keys, index=default_idx, key="line_palette")

    # Advanced Settings (Displayed directly to avoid Expander issues)
    st.markdown("#### 🎨 高级设置")
    c_a1, c_a2 = st.columns(2)
    # Use text_input as safe fallback for Stlite component issues
    lw_str = c_a1.text_input("线宽 (Line Width)", "2.0", key="line_lw")
    try: line_width = float(lw_str)
    except: line_width = 2.0
    
    ps_str = c_a2.text_input("点大小 (Point Size)", "6.0", key="line_ps")
    try: point_size = float(ps_str)
    except: point_size = 6.0
    
    c_a3, c_a4 = st.columns(2)
    show_reg = c_a3.checkbox("添加回归线 (Linear Fit for X-Numeric)", value=False, key="line_reg")
        # smooth = c_a4.checkbox("平滑曲线 (注: 仅视觉效果)", value=False, key="line_smooth")

    # Custom color
    custom_colors = None
    if palette_name == CUSTOM_PALETTE_KEY:
        custom_colors = st.text_input("输入自定义颜色 (逗号分隔)", placeholder="#FF0000, #00FF00", key="line_custom")

    if st.button("🚀 绘制折线图", type="primary", key="line_btn"):
        st.session_state['line_active'] = True

    if st.session_state.get('line_active', False):
        st.divider()
        nature_style.apply_nature_style()
        fig, ax = plt.subplots(figsize=(6, 5))

        # Data Prep
        plot_data = data.dropna(subset=[x_col, y_col]).copy()
        
        # Ensure X is sorted for line plot if numeric to avoid spaghetti plot
        is_numeric_x = pd.api.types.is_numeric_dtype(plot_data[x_col])
        if is_numeric_x:
             plot_data = plot_data.sort_values(by=x_col)

        hue = None if hue_col == "无" else hue_col
        
        # Determine N colors
        if hue:
            n_groups = plot_data[hue].nunique()
            colors = get_palette_colors(palette_name, n_colors=n_groups, custom_colors=custom_colors)
        else:
            colors = get_palette_colors(palette_name, n_colors=1, custom_colors=custom_colors)

        # Map error type to seaborn param
        # sns.lineplot errorbar: 'sd', 'se', ('ci', 95), None
        if error_type == "sd": err_param = "sd"
        elif error_type == "se": err_param = "se"
        elif "ci" in error_type: err_param = ("ci", 95)
        else: err_param = None

        # Main Plot
        sns.lineplot(data=plot_data, x=x_col, y=y_col, hue=hue, 
                     errorbar=err_param, palette=colors, ax=ax,
                     linewidth=line_width, marker="o" if show_points else None,
                     markersize=point_size)

        # Linear Regression
        if show_reg:
            if is_numeric_x:
                if hue:
                    st.warning("⚠️ 分组回归线在此简易模式下暂不支持叠加，仅绘制全局趋势作为参考。")
                
                # Draw global trend line
                sns.regplot(data=plot_data, x=x_col, y=y_col, scatter=False, ax=ax, 
                            color="gray", line_kws={"linestyle": "--", "alpha": 0.5})
            else:
                st.warning(f"⚠️ 无法绘制回归线：X轴 '{x_col}' 为非数值列 (文本/类别)。请选择时间或浓度等数值列作为 X 轴。")

        ax.set_xlabel(x_col)
        ax.set_ylabel(y_col)
        
        # Legend (Seaborn handles it, but we force position)
        if hue:
            plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)
        else:
            # Remove legend if no hue (regplot might add one if not careful, but here ok)
            pass

        sns.despine()
        st.pyplot(fig)
        
        # Stats info
        if is_numeric_x:
             try:
                 # Simple Regression Stats
                 res = stats.linregress(plot_data[x_col], plot_data[y_col])
                 st.info(f"📈 全局线性趋势 (Global Trend): R² = {res.rvalue**2:.3f}, p = {res.pvalue:.4e}")
             except: pass

        get_download_buttons(fig, f"Line_{y_col}", "line", report_title=f"Line Chart: {y_col} vs {x_col}")


# --- Density Plot Module ---
def render_density_module(data):
    st.header("📊 密度图 (Density/KDE Plot)")
    st.markdown("绘制核密度估计图，展示数据分布形态。")

    cols = data.select_dtypes(include=['number']).columns.tolist()
    all_cols = data.columns.tolist()
    
    if not cols:
        st.error("数据集中没有数值列")
        return

    c1, c2 = st.columns(2)
    val_col = c1.selectbox("数值列 (Value)", cols, key="kde_val")
    group_col = c2.selectbox("分组列 (Group, 可选)", ["无"] + all_cols, index=0, key="kde_group")

    c_opt1, c_opt2, c_opt3 = st.columns(3)
    fill_area = c_opt1.checkbox("填充区域 (Fill)", value=True, key="kde_fill")
    common_norm = c_opt2.checkbox("Common Norm", value=False, help="如果选中，所有组的总面积归一化为1；否则每组归一化为1（推荐）。")
    
    bw_str = c_opt3.text_input("平滑度 (Bandwidth, 0.1-2.0)", "1.0", key="kde_bw")
    try: bw_adjust = float(bw_str)
    except: bw_adjust = 1.0

    palette_keys = list(PALETTES.keys())
    default_idx = next((i for i, k in enumerate(palette_keys) if "Sci" in k), 0)
    palette_name = st.selectbox("配色方案", palette_keys, index=default_idx, key="kde_palette")

    # Custom color
    custom_colors = None
    if palette_name == CUSTOM_PALETTE_KEY:
        custom_colors = st.text_input("输入自定义颜色 (逗号分隔)", key="kde_custom")

    if st.button("🚀 绘制密度图", type="primary", key="kde_btn"):
        st.session_state['kde_active'] = True

    if st.session_state.get('kde_active', False):
        st.divider()
        nature_style.apply_nature_style()
        fig, ax = plt.subplots(figsize=(6, 5))
        
        plot_data = data.dropna(subset=[val_col])
        hue = None if group_col == "无" else group_col
        
        # Colors
        if hue:
            n_groups = plot_data[hue].nunique()
            colors = get_palette_colors(palette_name, n_colors=n_groups, custom_colors=custom_colors)
        else:
            colors = get_palette_colors(palette_name, n_colors=1, custom_colors=custom_colors)

        sns.kdeplot(data=plot_data, x=val_col, hue=hue, 
                    fill=fill_area, common_norm=common_norm, bw_adjust=bw_adjust,
                    palette=colors, alpha=0.5, linewidth=1.5, ax=ax)

        ax.set_xlabel(val_col)
        sns.despine()
        
        if hue:
            plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)

        st.pyplot(fig)
        get_download_buttons(fig, f"KDE_{val_col}", "kde", report_title=f"Density Plot: {val_col}")

def render_survival_module(data):
    st.header("💀 生存分析 (Survival Analysis - KM)")
    cols = data.columns.tolist()
    
    c1, c2, c3 = st.columns(3)
    time_col = c1.selectbox("时间列 (Time)", cols, key="surv_time")
    event_col = c2.selectbox("事件列 (Event, 1=Dead)", cols, key="surv_event")
    group_col = c3.selectbox("分组列 (Group)", cols, key="surv_group")
    
    palette_name = st.selectbox("配色方案", list(PALETTES.keys()), index=0, key="surv_palette")
    
    if st.button("🚀 绘制生存曲线", type="primary"):
        st.session_state['surv_active'] = True

    if st.session_state.get('surv_active', False):
        st.divider()
        nature_style.apply_nature_style()
        fig, ax = plt.subplots(figsize=(6, 6))
        
        kmf = KaplanMeierFitter()
        groups = sorted(data[group_col].unique())
        colors = get_palette_colors(palette_name, n_colors=len(groups))
        
        results_text = []
        
        for i, g in enumerate(groups):
            mask = data[group_col] == g
            kmf.fit(data.loc[mask, time_col], data.loc[mask, event_col], label=str(g))
            kmf.plot_survival_function(ax=ax, ci_show=False, linewidth=2, color=colors[i])
            
        # Log-rank test
        if len(groups) == 2:
            g1, g2 = groups[0], groups[1]
            try:
                res = logrank_test(
                    data[data[group_col]==g1][time_col], data[data[group_col]==g2][time_col],
                    event_observed_A=data[data[group_col]==g1][event_col], event_observed_B=data[data[group_col]==g2][event_col]
                )
                p_val = res.p_value
                ax.text(0.05, 0.1, f"Log-rank p = {p_val:.4f}", transform=ax.transAxes, fontsize=10)
                results_text.append(f"Log-rank test ({g1} vs {g2}): p={p_val:.4f}")
            except Exception as e:
                st.write(f"Log-rank Warning: {e}")

        ax.set_xlabel("Time")
        ax.set_ylabel("Survival Probability")
        ax.set_ylim(0, 1.05)
        sns.despine()
        st.pyplot(fig)
        
        get_download_buttons(fig, "Survival_Curve", "surv", report_title="Kaplan-Meier Survival Analysis")

def render_heatmap_module(data):
    st.header("🔥 热图聚类 (Heatmap)")
    num_cols = data.select_dtypes(include=['number']).columns.tolist()
    all_cols = data.columns.tolist()
    
    cols = st.multiselect("数据列", num_cols, default=num_cols[:10])
    idx = st.selectbox("索引列 (Row Label)", ["Auto-Index"] + all_cols)
    
    if st.button("🚀 绘制热图", type="primary"):
        st.session_state['heat_active'] = True

    if st.session_state.get('heat_active', False):
        df_heat = data[cols].copy()
        if idx != "Auto-Index": df_heat.index = data[idx]
        df_heat = df_heat.dropna()
        
        nature_style.apply_nature_style()
        g = sns.clustermap(df_heat, z_score=0, cmap="vlag", center=0, figsize=(6,6))
        st.pyplot(g)
        get_download_buttons(g, "Heatmap", "heat", report_title="Hierarchical Clustering Heatmap")

def render_pca_module(data):
    st.header("🧬 PCA 分析 (3D Supported)")
    num_cols = data.select_dtypes(include=['number']).columns.tolist()
    feats = st.multiselect("特征列", num_cols, default=num_cols[:5])
    
    c1, c2 = st.columns(2)
    label = c1.selectbox("着色", data.columns)
    palette_name = c2.selectbox("配色方案", list(PALETTES.keys()), index=0, key="pca_palette")
    
    use_3d = st.checkbox("3D 模式")
    
    if st.button("🚀 运行 PCA"):
        st.session_state['pca_active'] = True

    if st.session_state.get('pca_active', False):
        df_pca = data[feats].dropna()
        scaler = StandardScaler()
        pcs = PCA(n_components=3).fit_transform(scaler.fit_transform(df_pca))
        pca_df = pd.DataFrame(pcs, columns=['PC1','PC2','PC3'])
        
        # Ensure label matches index of cleaned data
        labels_aligned = data.loc[df_pca.index, label].values
        pca_df['Label'] = labels_aligned
        
        unique_labels = sorted(list(set(labels_aligned)))
        colors = get_palette_colors(palette_name, n_colors=len(unique_labels))
        
        nature_style.apply_nature_style()
        fig = plt.figure(figsize=(6,6))
        
        if use_3d:
            ax = fig.add_subplot(111, projection='3d')
            # Manually map colors for 3D scatter
            color_map = {lbl: col for lbl, col in zip(unique_labels, colors)}
            c_array = [color_map[l] for l in pca_df['Label']]
            
            sc = ax.scatter(pca_df.PC1, pca_df.PC2, pca_df.PC3, c=c_array, s=50)
            
            # Create manual legend
            import matplotlib.patches as mpatches
            patches = [mpatches.Patch(color=color_map[l], label=l) for l in unique_labels]
            ax.legend(handles=patches)
            
        else:
            sns.scatterplot(data=pca_df, x='PC1', y='PC2', hue='Label', palette=colors)
            
        st.pyplot(fig)
        get_download_buttons(fig, "PCA", "pca", report_title="PCA Analysis")

# --- Report Generation Page ---
def render_report_page():
    st.title("📝 实验报告生成 (Generator)")
    
    if not st.session_state['report_items']:
        st.info("购物车是空的。请在分析模块中点击 '➕ 添加到报告'。")
        return

    st.write(f"当前共有 {len(st.session_state['report_items'])} 个项目。")
    
    # Preview
    for i, item in enumerate(st.session_state['report_items']):
        with st.expander(f"{i+1}. {item['title']} ({item['time']})"):
            if item['text']: st.write(item['text'])
            if item['img_bytes']: st.image(item['img_bytes'])
            if item['df'] is not None: st.dataframe(item['df'])
            if st.button("❌ 删除", key=f"del_{i}"):
                st.session_state['report_items'].pop(i)
                st.rerun()

    if st.button("💾 生成 Word 报告 (.docx)", type="primary"):
        doc = Document()
        doc.add_heading('Bio-Analysis Experiment Report', 0)
        doc.add_paragraph(f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        for item in st.session_state['report_items']:
            doc.add_heading(item['title'], level=1)
            
            if item['text']:
                doc.add_paragraph(item['text'])
            
            if item['img_bytes']:
                try:
                    doc.add_picture(io.BytesIO(item['img_bytes']), width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"[Image Error: {e}]")
            
            if item['df'] is not None:
                # Add Table
                df = item['df']
                table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                table.style = 'Table Grid'
                
                # Header
                for j, col_name in enumerate(df.columns):
                    table.cell(0, j).text = str(col_name)
                
                # Rows
                for r in range(df.shape[0]):
                    for c in range(df.shape[1]):
                        table.cell(r+1, c).text = str(df.iloc[r, c])
                doc.add_paragraph("") # Space
        
        # Save
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label="📥 点击下载 Word 报告",
            data=bio.getvalue(),
            file_name=f"Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- Data Format Guide ---
def render_data_guide():
    with st.expander("📌 数据格式指南 (Data Format Guide)", expanded=True):
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("### 1. 长格式 (Long Format)")
            st.info("适用于：差异分析、生存分析 (Diff/Survival)")
            st.dataframe(pd.DataFrame({
                "Group": ["Ctrl", "Ctrl", "Treat", "Treat"],
                "Weight": [20.1, 20.5, 25.3, 24.8],
                "Time": [10, 12, 10, 11]
            }), height=150)
            
        with c2:
            st.markdown("### 2. 宽格式 (Wide Format)")
            st.info("适用于：PCA、热图 (PCA/Heatmap)")
            st.dataframe(pd.DataFrame({
                "Sample": ["S1", "S2", "S3"],
                "GeneA": [1.2, 5.6, 1.1],
                "GeneB": [3.4, 1.2, 3.5],
                "Group": ["WT", "KO", "WT"]
            }), height=150)

# --- Main App Structure ---
st.sidebar.title("控制面板")
mode = st.sidebar.radio("功能模块", [
    "🏠 首页 & 指南",
    "📊 箱线图 (Boxplot)",
    "📊 条形图",
    "📈 折线图 (Line)",
    "📊 密度图 (Density)",
    "💀 生存分析",
    "🧬 PCA (3D)",
    "🔥 热图聚类",
    "📈 相关性",
    "📝 导出报告"
])

if mode == "🏠 首页 & 指南":
    st.title("💎 ElementPrism Workbench")
    st.markdown("欢迎使用 ElementPrism 生物数据可视化分析套件。请在左侧上传数据并选择模块。")
    render_data_guide()
    
elif mode == "📝 导出报告":
    render_report_page()
    
else:
    uploaded_file = st.sidebar.file_uploader("📂 上传 Excel/CSV 数据", type=['xlsx', 'csv'])
    if uploaded_file:
        data = load_data(uploaded_file)
        if data is not None:
            # Use the top-level 'mode' variable directly to cleaner logic
            if "描述" in mode: render_desc_stats(data)
            elif "箱线图" in mode: render_difference_module(data)
            elif "条形图" in mode: render_barplot_module(data)
            elif "折线图" in mode: render_linechart_module(data)
            elif "密度图" in mode: render_density_module(data)
            elif "PCA" in mode: render_pca_module(data)
            elif "热图" in mode: render_heatmap_module(data)
            elif "生存" in mode: render_survival_module(data)
            elif "相关" in mode: 
                # Re-implement simple correlation for v4
                st.header("📈 相关性分析")
                num_cols = data.select_dtypes(include=['number']).columns
                x = st.selectbox("X", num_cols, key="cx")
                y = st.selectbox("Y", num_cols, key="cy")
                
                col_w1, col_w2 = st.columns(2)
                fig_width = col_w1.slider("图表宽度 (inches)", 3, 15, 6)
                fig_height = col_w2.slider("图表高度 (inches)", 3, 15, 6)

                if st.button("Run"):
                    st.session_state['corr_active'] = True

                if st.session_state.get('corr_active', False):
                    nature_style.apply_nature_style()
                    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
                    sns.regplot(data=data, x=x, y=y, ax=ax)
                    st.pyplot(fig)
                    get_download_buttons(fig, "Corr", "corr", report_title=f"Correlation {x} vs {y}")
    else:
        st.info("👈 请先上传数据")
